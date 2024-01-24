#!/usr/bin/env python3
# -*- coding: utf-8 -*-
""" Write document file. """

import re
from typing import List, Union

import docx
from docx.shared import Inches

from dessia_common.files import BinaryFile


class LayoutElement:
    """
    Represents a header or footer in document.

    This class is used to create headers and footers for document. Subclasses
    `Header` and `Footer` are available for convenience.
    """

    def __init__(self, text: str, align: str = 'center'):
        self.text = text
        self.align = align

    def _add_to_section(self, section, type_: str):
        """
        Add the header or footer to the specified section.

        :param section: The `docx.section.Section` object to add the header/footer to.
        :param type_: The type of element to add, either "header" or "footer".
        """
        element = getattr(section, type_)
        paragraph = element.add_paragraph()
        paragraph.text = self.text
        paragraph.alignment = getattr(docx.enum.text.WD_PARAGRAPH_ALIGNMENT, self.align.upper())

    def _add_picture(self, section, type_: str, image_path: str, width: int = None, height: int = None):
        """ Add the header or footer picture to the specified section. """
        element = getattr(section, type_)
        paragraph = element.paragraphs[0] if element.paragraphs else element.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=width, height=height)


class Header(LayoutElement):
    """ Represents a header in document. """

    def add_to_document(self, document: docx.Document):
        """ Add the header to the document. """
        section = document.sections[0]
        super()._add_to_section(section=section, type_='header')

    def add_to_section(self, section):
        """ Add header to section. """
        super()._add_to_section(section=section, type_='header')

    def add_picture(self, section, image_path: str, width: float = None, height: float = None):
        """ Add picture for header. """
        if width:
            width = Inches(width)
        if height:
            height = Inches(height)
        super()._add_picture(section=section, type_='header', image_path=image_path, width=width, height=height)


class Footer(LayoutElement):
    """ Represents a footer in document. """

    def add_to_document(self, document: docx.Document):
        """ Add the footer to the document. """
        section = document.sections[0]
        super()._add_to_section(section=section, type_='footer')

    def add_to_section(self, section):
        """ Add footer to section. """
        super()._add_to_section(section=section, type_='footer')

    def add_picture(self, section, image_path: str, width: float = None, height: float = None):
        """ Add picture for footer. """
        if width:
            width = Inches(width)
        if height:
            height = Inches(height)
        super()._add_picture(section=section, type_='footer', image_path=image_path, width=width, height=height)


class Heading:
    """ Represents a heading in document. """

    def __init__(self, text: str, level: int):
        self.text = text
        self.level = level

    def add_to_document(self, document: docx.Document):
        """ Add the heading to the document. """
        document.add_heading(self.text, self.level)

    @classmethod
    def from_markdown(cls, markdown_text: str):
        """ Create a new Paragraph object from Markdown text. """
        return cls(text=markdown_text.strip("#"), level=markdown_text.count("#"))


class Paragraph:
    """ Represents a paragraph in the document. """

    def __init__(self, text: str):
        self.text = text

    def add_to_document(self, document: docx.Document):
        """ Add paragraph to the document. """
        document.add_paragraph(self.text)

    @classmethod
    def from_markdown(cls, markdown_text: str):
        """ Create a new Paragraph object from Markdown text. """
        return cls(text=markdown_text)


class Table:
    """ Represents a table in the document. """

    def __init__(self, rows: List[List[str]]):
        self.rows = rows

    def add_to_document(self, document: docx.Document, style: str = "TableGrid"):
        """
        Add the table to the document.

        :param document: The document to which the table will be added.
        :param style: The table style to use. Default is "TableGrid". Other available styles include:
                        - "LightShading": adds light shading to the table cells.
                        - "LightGrid": similar to "TableGrid", but with lighter grid lines.
                        - "MediumShading1", "MediumShading2": adds medium-level shading to the table cells,
                        with different levels of intensity.
                        - "DarkList": displays the table as a numbered list, with a dark background color for each item.
                        - "ColorfulGrid": uses bright, contrasting colors for the table borders and cell backgrounds.
                        - "LightList": displays the table as a bulleted list, with a light background color for each
                        item.
        """
        table = document.add_table(rows=len(self.rows), cols=len(self.rows[0]), style=style)
        for i, row in enumerate(self.rows):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell


class Section:
    """
    Represents a section in the document.

    A section is a part of the document with its own headers, footers.
    """

    def __init__(self):
        self.elements = []

    def add_element(self, element: Union[Header, Footer]):
        """ Add a header to the section. """
        self.elements.append(element)

    def add_to_document(self, document: docx.Document):
        """ Add the section to the document. """
        section = document.sections[0]
        for element in self.elements:
            element.add_to_section(section)

    def add_picture_to_document(self, document: docx.Document, image_path, width, height):
        """ Add picture to the section. """
        section = document.sections[0]
        if self.elements:
            self.elements[0].add_picture(section, image_path, width, height)


class DocxWriter:
    """ Write a Document file. """

    def __init__(self, paragraphs: List[Paragraph] = None, section: Section = None, filename: str = "document.docx",
                 headings: List[Heading] = None, tables: List[Table] = None):
        self.filename = filename
        if headings is None:
            headings = []
        self.headings = headings
        self.section = section
        if paragraphs is None:
            paragraphs = []
        self.paragraphs = paragraphs
        self.tables = tables
        self.document = docx.Document()

        if self.section:
            self.section.add_to_document(document=self.document)

    def add_headings(self, page_break: bool = True) -> 'DocxWriter':
        """ Add a list of headings to the document. """
        for heading in self.headings:
            self.document.add_heading(heading.text, level=heading.level)
        if page_break:
            self.add_page_breaks(num_page_breaks=1)
        return self

    def add_paragraphs(self, add_heading: bool = True) -> 'DocxWriter':
        """ Add a list of paragraphs to the document. """
        for paragraph in self.paragraphs:
            if self.headings and add_heading:
                self.document.add_heading(self.headings[0].text, self.headings[0].level)
                self.headings = self.headings[1:]
            self.document.add_paragraph(paragraph.text)
        return self

    def add_paragraph_as_heading(self, text: str):
        """
        Adds a new heading to the document, using the specified text as the heading text.

        The level of the heading is determined by the number of '#' characters in the text.
        For example, a single '#' character at the beginning of the text will create a level 1 heading.
        """
        self.document.add_heading(text, level=text.count('#'))
        return self

    def add_page_breaks(self, num_page_breaks: int):
        """ Add an empty page to the document. """
        for _ in range(num_page_breaks):
            self.document.add_page_break()

    def add_table(self, add_all_tables: bool = False, table_index: int = 0) -> 'DocxWriter':
        """ Add tables to the document. """
        if add_all_tables:
            for table in self.tables:
                table.add_to_document(self.document)
                self.document.add_paragraph()
        else:
            self.tables[table_index].add_to_document(self.document)
        return self

    def add_list_items(self, items: List[str], style: str = 'List Bullet'):
        """
        Add a list (bullet or numbered) to the document.

        :param items: A list of strings to be added as list items
        :type items: List[str]

        :param style: The style of the list. Valid values are 'List Bullet' (default) or 'List Number'
        :type style: str
        """
        for item in items:
            self.document.add_paragraph(item, style=style)
        return self

    def add_header_footer_picture(self, image_path: str, width: int = None, height: int = None)\
            -> 'DocxWriter':
        """ Add a picture to the header or footer of the document. """
        self.section.add_picture_to_document(document=self.document, image_path=image_path, width=width, height=height)
        return self

    def add_picture(self, image_path: str, width: float = None, height: float = None) -> 'DocxWriter':
        """ Add an image to the document. """
        if width:
            width = Inches(width)
        if height:
            height = Inches(height)
        self.document.add_picture(image_path, width=width, height=height)
        return self

    def delete_layout_element(self):
        """ Remove the header and footer from all sections in a Word document. """
        for section in self.document.sections:
            section.different_first_page_header_footer = False
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True

    @staticmethod
    def parse_markdown(markdown_text: str):
        """
        Parses the given markdown text.

        :return: Tuple containing a list of headings, a list of paragraphs and tables.
        """
        elements, headings = [], []
        table_pattern = re.compile(r'^\|.*\|$')
        horizontal_line_pattern = re.compile(r'^\s*\|?\s*-+\s*\|?\s*(-+\s*\|?)*\s*$')

        for line in markdown_text.split('\n'):
            line = line.strip()
            if line:
                if line.startswith('#'):
                    headings.append(Heading.from_markdown(line))
                    elements.append(Paragraph.from_markdown(markdown_text=line))

                elif table_pattern.match(line) and not horizontal_line_pattern.match(line):
                    row = line.strip('|').split('|')
                    if '---' not in line:
                        elements.append(Table(rows=[row]))

                else:
                    if '---' not in line:
                        elements.append(Paragraph.from_markdown(markdown_text=line))
        return headings, elements

    @classmethod
    def from_markdown(cls, markdown_text: str):
        """ Converts the given markdown text into a DocxWriter instance."""
        headings, elements = cls.parse_markdown(markdown_text=markdown_text)
        docx_writer = cls(headings=headings)
        docx_writer.add_headings()
        for item in elements:
            if isinstance(item, Paragraph):
                docx_writer.paragraphs = [item]
                if item.text.startswith('#'):
                    docx_writer.add_paragraph_as_heading(text=item.text.strip(' # '))
                else:
                    docx_writer.add_paragraphs(add_heading=False)
            if isinstance(item, Table):
                docx_writer.tables = [item]
                docx_writer.add_table()
        return docx_writer

    def save_file(self):
        """ Saves the document to a file. """
        if not self.filename.endswith('.docx'):
            self.filename += '.docx'
        self.document.save(self.filename)

    def save_to_stream(self, stream: BinaryFile):
        """ Saves the document to a binary stream. """
        self.document.save(stream)
